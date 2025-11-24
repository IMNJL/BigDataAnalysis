"""Generate a DOCX report with theoretical notes, experiment results, a flowchart, and the provided mind map image.

Usage:
    python generate_docx.py --out outputs/AutoGPT_experiment_results.docx
"""
import os
import json
import argparse
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from PIL import Image

HERE = os.path.dirname(__file__)
AUTO_STATS = os.path.join(HERE, 'outputs', 'auto', 'stats.json')
MANUAL_STATS = os.path.join(HERE, 'outputs', 'manual', 'stats.json')
COMPARE = os.path.join(HERE, 'outputs', 'comparison_summary.md')
PREFERRED_FLOW_IMAGES = [os.path.join(HERE, 'image copy.png'), os.path.join(HERE, 'workflow.png')]
FLOWCHART_PNG = os.path.join(HERE, 'outputs', 'flowchart.png')
for p in PREFERRED_FLOW_IMAGES:
    if os.path.exists(p):
        FLOWCHART_PNG = p
        break

MINDMAP_PNG = os.path.join(HERE, 'image.png')


def load_json(path):
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as fh:
        return json.load(fh)


def load_text(path):
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as fh:
        return fh.read()


def draw_flowchart(path):
    # Simple matplotlib flowchart: three boxes with arrows
    fig, ax = plt.subplots(figsize=(8,3))
    ax.axis('off')

    # Boxes positions
    box_props = dict(boxstyle='round,pad=0.6', fc='#f0f0f0', ec='black')

    ax.text(0.15, 0.5, 'Task Generation', ha='center', va='center', bbox=box_props, fontsize=12)
    ax.text(0.5, 0.5, 'Task Execution', ha='center', va='center', bbox=box_props, fontsize=12)
    ax.text(0.85, 0.5, 'Self-Reflection', ha='center', va='center', bbox=box_props, fontsize=12)

    # Arrows
    ax.annotate('', xy=(0.3, 0.5), xytext=(0.36, 0.5), arrowprops=dict(arrowstyle='->', lw=2))
    ax.annotate('', xy=(0.65, 0.5), xytext=(0.71, 0.5), arrowprops=dict(arrowstyle='->', lw=2))

    # Sub-notes under each box
    ax.text(0.15, 0.28, 'decompose goal\ncreate tasks\nprioritize', ha='center', va='center', fontsize=9)
    ax.text(0.5, 0.28, 'execute tools\nrun code\ncollect artifacts', ha='center', va='center', fontsize=9)
    ax.text(0.85, 0.28, 'evaluate results\nreplan / refine\nstore memory', ha='center', va='center', fontsize=9)

    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def build_docx(out_path):
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    auto = load_json(AUTO_STATS)
    manual = load_json(MANUAL_STATS)
    comp = load_text(COMPARE)

    doc = Document()
    doc.add_heading('AutoGPT Agent Experiment Report', level=1)

    # 1) Design the workflow of an AI Agent (use provided workflow image if available)
    doc.add_heading('1. Design: AI Agent Workflow', level=2)
    doc.add_paragraph('The workflow diagram below shows the experiment orchestration steps for running an AutoGPT-style data analysis task.')
    if os.path.exists(FLOWCHART_PNG):
        doc.add_picture(FLOWCHART_PNG, width=Inches(6))
    else:
        # fallback to generated simple flowchart
        draw_flowchart(FLOWCHART_PNG)
        doc.add_picture(FLOWCHART_PNG, width=Inches(6))

    doc.add_paragraph('Step descriptions:')
    doc.add_paragraph('1. Environment Preparation: install Docker/Git/Python and configure environment variables (e.g., OpenAI API key).')
    doc.add_paragraph('2. System Setup: run the startup script (for AutoGPT or your agent) and ensure dependencies are available.')
    doc.add_paragraph('3. Input Prompt: provide the task prompt (e.g., "Analyze this dataset and generate a report") and parse the task into a plan; the agent may wait for user confirmation.')
    doc.add_paragraph('4. Execute the Think-Act-Observe Cycle: load the dataset, inspect structure, calculate descriptive statistics, create visualizations, and generate report artifacts.')
    doc.add_paragraph('5. Output the Report: collect generated artifacts and output the final report (markdown, docx, or PDF).')

    # 2) Mind map
    doc.add_heading('2. Mind map (lesson content)', level=2)
    doc.add_paragraph('High-level components: agent, planning, tools, memory, and reflection.')
    if os.path.exists(MINDMAP_PNG):
        img = Image.open(MINDMAP_PNG)
        max_w = 1600
        if img.width > max_w:
            ratio = max_w / img.width
            img = img.resize((int(img.width*ratio), int(img.height*ratio)))
            resized_path = os.path.join(HERE, 'outputs', 'mindmap_resized.png')
            img.save(resized_path)
            doc.add_picture(resized_path, width=Inches(6))
        else:
            doc.add_picture(MINDMAP_PNG, width=Inches(6))
    else:
        doc.add_paragraph('Mind map image not found: image.png')

    # 3) Comparison: Agents vs traditional LLMs
    doc.add_heading('3. Comparison: Agents vs traditional LLMs', level=2)
    doc.add_paragraph('Traditional LLMs (single-turn or chat models) generate text given a prompt. They excel at transformation, summarization, and single-step tasks but cannot autonomously plan, call tools, or maintain long-running state without orchestration.')
    doc.add_paragraph('AI Agents extend LLMs with:')
    doc.add_paragraph('- Autonomous task planning and decomposition', style='List Bullet')
    doc.add_paragraph('- Tool use (code execution, web search, calculators)', style='List Bullet')
    doc.add_paragraph('- Memory and stateful operation across steps', style='List Bullet')
    doc.add_paragraph('- Self-reflection and iterative improvement', style='List Bullet')

    # 4) Experiment: Using AutoGPT (results)
    doc.add_heading('4. Experiment: Using AutoGPT', level=2)
    doc.add_paragraph('Experiment 1: Have AutoGPT complete a data analysis report on the Titanic dataset. Observe planning (task list), execution (artifacts), and reflection (replanning).')

    doc.add_paragraph('Auto agent detected columns and statistics:')
    if auto:
        for k, v in auto.items():
            doc.add_paragraph(f'- {k}: {v}', style='List Bullet')
    else:
        doc.add_paragraph('Auto agent stats not found.')

    doc.add_paragraph('Manual agent detected columns and statistics:')
    if manual:
        for k, v in manual.items():
            doc.add_paragraph(f'- {k}: {v}', style='List Bullet')
    else:
        doc.add_paragraph('Manual agent stats not found.')

    doc.add_paragraph('Comparison summary:')
    if comp:
        doc.add_paragraph(comp)
    else:
        doc.add_paragraph('Comparison summary not found.')

    doc.add_heading('5. Notes & next steps', level=2)
    doc.add_paragraph('Potential extensions: attach full logs, include side-by-side visual comparisons of plots, or integrate an actual AutoGPT framework and record its internal plan outputs for richer analysis.')

    doc.add_page_break()
    doc.save(out_path)
    print('Saved DOCX to', out_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--out', default=os.path.join(HERE, 'outputs', 'AutoGPT_experiment_results.docx'))
    args = parser.parse_args()
    build_docx(args.out)
