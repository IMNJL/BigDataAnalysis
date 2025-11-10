"""
Simple report generator for the BDA_3 AI-enhanced ETL experiment.
Generates a DOCX file with the required template sections and inserts any available
processed JSON results and screenshots from `screenshots/`.
"""
import os
import argparse
from docx import Document
from docx.shared import Pt, Inches
import json

ROOT = os.path.dirname(__file__)
PROCESSED_DIR = os.path.join(ROOT, 'data', 'processed')
SCREENSHOT_DIR = os.path.join(ROOT, 'screenshots')
OUT_DOC = os.path.join(ROOT, 'BDA_3_ETL_report.docx')


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)


def insert_results(doc):
    if not os.path.exists(PROCESSED_DIR):
        add_paragraph(doc, 'No processed data found. Run cleaner.py first to generate processed JSON.')
        return
    files = sorted([f for f in os.listdir(PROCESSED_DIR) if f.endswith('.json')])
    if not files:
        add_paragraph(doc, 'No processed JSON files found in data/processed/')
        return
    add_heading(doc, 'Processed files and extracted comments', level=2)
    for fn in files:
        path = os.path.join(PROCESSED_DIR, fn)
        try:
            with open(path, 'r', encoding='utf-8') as f:
                j = json.load(f)
            add_paragraph(doc, f"Source: {j.get('source')}")
            add_paragraph(doc, f"Full text length: {j.get('full_text_length')}")
            add_paragraph(doc, 'Top comments:')
            for c in j.get('comments', [])[:5]:
                add_paragraph(doc, f"- {c}")
            doc.add_page_break()
        except Exception as e:
            add_paragraph(doc, f"Failed to read {fn}: {e}")


def insert_screenshots_and_diagrams(doc):
    """Insert architecture diagram (architecture.png) if present and any screenshots from SCREENSHOT_DIR."""
    # insert architecture diagram first if exists
    arch_path = os.path.join(SCREENSHOT_DIR, 'architecture.png')
    if os.path.exists(arch_path):
        add_heading(doc, 'Architecture Diagram', level=2)
        try:
            doc.add_picture(arch_path, width=Inches(6))
            add_paragraph(doc, 'Figure: Architecture diagram')
        except Exception as e:
            add_paragraph(doc, f'Failed to insert architecture diagram: {e}')

    # insert any other screenshot images
    imgs = []
    if os.path.exists(SCREENSHOT_DIR):
        for fn in sorted(os.listdir(SCREENSHOT_DIR)):
            if fn.lower().endswith(('.png', '.jpg', '.jpeg')) and fn != 'architecture.png':
                imgs.append(os.path.join(SCREENSHOT_DIR, fn))
    if imgs:
        add_heading(doc, 'Screenshots', level=2)
        # try to read captions from captions.json (optional)
        captions_path = os.path.join(SCREENSHOT_DIR, 'captions.json')
        captions = {}
        try:
            if os.path.exists(captions_path):
                import json
                with open(captions_path, 'r', encoding='utf-8') as cf:
                    captions = json.load(cf)
        except Exception:
            captions = {}

        for p in imgs:
            try:
                doc.add_picture(p, width=Inches(5.5))
                base = os.path.basename(p)
                cap = captions.get(base) if isinstance(captions, dict) else None
                if cap:
                    add_paragraph(doc, f'Screenshot: {cap}')
                else:
                    add_paragraph(doc, f'Screenshot: {base}')
            except Exception as e:
                add_paragraph(doc, f'Failed to insert screenshot {p}: {e}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--experiment-name', '-n', default='BDA_3 AI-Enhanced ETL')
    args = parser.parse_args()

    doc = Document()
    add_heading(doc, args.experiment_name)

    add_heading(doc, 'Part 1: Experiment Plan', level=2)
    add_paragraph(doc, 'Goal: extract data and related comment information from PDF files using an AI-enhanced ETL method.')
    add_paragraph(doc, 'Inputs: Topic or seed URL(s) or direct PDF links.')
    add_paragraph(doc, 'Outputs: cleaned JSON per PDF and this experiment report.')

    add_heading(doc, 'Part 2: Architecture', level=2)
    add_paragraph(doc, 'Architecture: simple pipeline: Crawler -> PDF download -> Extract -> Cleaner (heuristics/LLM) -> Report generator')
    add_paragraph(doc, 'Insert your architecture diagram as screenshots/architecture.png to include it here.')

    add_heading(doc, 'Part 3: Steps and Screenshots', level=2)
    add_paragraph(doc, 'Step 1: Run crawler to download PDFs into data/raw_pdfs\nPrompt: (seed URL)\nResponse: PDF files downloaded into data/raw_pdfs')
    add_paragraph(doc, 'Step 2: Run cleaner to extract text and comments from PDFs\nPrompt: (run cleaner)\nResponse: JSON files in data/processed')
    add_paragraph(doc, 'Add screenshots into screenshots/ and they will be inserted here in order.')

    insert_results(doc)

    doc.save(OUT_DOC)
    print(f"Wrote report to {OUT_DOC}")
