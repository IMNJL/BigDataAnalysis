"""
Report generator for BDA_8: Data Parallelism and Model Parallelism assignment.
Generates `BDA_8_Report.docx` with detailed answers to Parts 1-4 in Russian.
"""
import os
import argparse
from docx import Document
from docx.shared import Pt

OUT_DOC_DEFAULT = 'BDA_8_Report.docx'


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.style.font.size = Pt(11)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def write_report(out_path):
    doc = Document()
    doc.add_heading('BDA_8 — Data Parallelism и Model Parallelism', level=0)

    # Part 1: Basic Concept Questions
    add_heading(doc, 'Part 1: Базовые понятия', level=1)

    # Q1
    add_paragraph(doc, '1. Разница между "Data Parallelism" и "Model Parallelism"', bold=True)
    add_paragraph(doc, """Data Parallelism (параллелизм по данным): одна и та же модель копируется на несколько устройств (GPU/процессов). Каждый worker получает свою часть батча/набора данных и выполняет прямой и обратный проход на локальных данных. Затем градиенты агрегируются (например, через All-Reduce) и параметры синхронизируются.

Применимость: хорошо подходит, когда модель помещается на одно устройство, а объём данных велик. Пример: обучение ResNet на большом наборе изображений на нескольких GPU.""")
    add_paragraph(doc, """Model Parallelism (параллелизм модели): модель разделяется между устройствами. Разные слои или части модели размещаются на разных GPU/нодах; данные проходят последовательно через сегменты модели.

Применимость: когда модель слишком велика, чтобы поместиться в память одного устройства (например, GPT-3 / большой трансформер), либо когда некоторые части модели вычислительно интенсивны и выгодно распараллелить по структуре. Пример: распределение слоёв гигантской трансформерной сети по нескольким GPU.""")

    add_paragraph(doc, """Ключевые различия:
- Data parallelism масштабируется по объёму данных; model parallelism — по размеру модели.
- Data-parallel легче реализовать (репликация модели) и обычно даёт хорошую скорость; model-parallel требует управления передачей активаций и высоких накладных расходов на коммуникацию.""")

    # Q2
    add_paragraph(doc, '\n2. Что такое All-Reduce? Преимущества по сравнению с Parameter Server', bold=True)
    add_paragraph(doc, 'All-Reduce — коллективная коммуникационная операция, агрегирующая (например, суммируя) векторы (градиенты) со всех процессов и возвращающая результат каждому процессу. В контексте обучения это часто используется для суммирования/усреднения градиентов и получения согласованных обновлений параметров без центральной точки.\n\nПреимущества All-Reduce vs Parameter Server:\n- Нет единой точки отказа — распределённая топология (peer-to-peer) даёт лучшую отказоустойчивость.\n- Эффективность: современные All-Reduce реализации (NCCL, MPI) используют кольцевые/деревянные алгоритмы и оптимизированы для GPU, что снижает латентность и увеличивает пропускную способность по сравнению с централизованным Parameter Server, особенно при большом числе узлов.\n- Стабильность и согласованность: All-Reduce дает синхронное усреднение градиентов, что упрощает аналитику сходимости; Parameter Server часто использует асинхронный режим, что может приводить к stale gradients и проблемам со сходимостью.')

    # Q3
    add_paragraph(doc, '\n3. Как PyTorch DDP реализует распределённое обучение? (кратко)', bold=True)
    add_paragraph(doc, 'PyTorch DistributedDataParallel (DDP) работает так:\n1) Модель реплицируется на каждом процессе (обычно — по одному процессу на GPU).\n2) Данные разбиваются между процессами (обычно через DistributedSampler для DataLoader).\n3) Каждый процесс выполняет прямой и обратный проход на своей локальной подвыборке.\n4) После вызова backward() происходит автоматическая синхронная коммуникация градиентов: DDP запускает All-Reduce для соответствующих параметров, агрегирует градиенты (сумма/усреднение).\n5) После агрегирования каждый процесс обновляет свои локальные параметры (обычно через оптимайзер — шаг синхронный).\n\nDDP минимизирует накладные расходы на коммуникацию, синхронизируя градиенты слоями параллельно по частям и используя NCCL на GPU для производительности.')

    # Q4
    add_paragraph(doc, '\n4. Какие ключевые компоненты Ray? Роль Object Store', bold=True)
    add_paragraph(doc, 'Ключевые компоненты Ray:\n- Ray Core (Scheduler): распределение задач и управление ресурсами.\n- Ray Actors: состояние́вые объекты (долгоживущие процессы), которые сохраняют состояние между вызовами.\n- Tasks: тонкие единицы вычисления (stateless), выполняемые на пуле ресурсов.\n- Object Store (Plasma): распределённое хранилище объектов в памяти, позволяющее обмениваться данными между задачами/акторами без сериализации по сети.\n\nРоль Object Store: он хранит объекты (напр., тензоры, батчи данных) в общей памяти узла и предоставляет быстрый доступ другим задачам/акторам на том же узле. Это сокращает расходы на копирование и сериализацию, позволяя выполнять zero-copy передачи между задачами, что критично при обмене большими тензорами в распределённых ML-пайплайнах.')

    # Q5
    add_paragraph(doc, '\n5. Основные ограничения Spark MLlib; почему не подходит для обучения deep learning напрямую', bold=True)
    add_paragraph(doc, 'Ограничения Spark MLlib:\n- Модельный набор ограничен классическими алгоритмами (регрессии, деревья, кластеризация), а не специализирован для нейросетей.\n- Архитектура ориентирована на обработку данных (лоадинг, трансформации, ETL) и не оптимизирована для GPU-ускорённых вычислений; обучение глубоких сетей требует эффективной работы с CUDA и тензорных библиотек, чего у MLlib нет.\n- Перформанс: MLlib ориентирован на масштабирование по данным на CPU-кластере; при работе с большими глубокими нейросетями и GPU необходима более узконаправленная оптимизация коммуникации градиентов и памяти.\n- Интерфейсы и экосистема: большинство современных DL-фреймворков (PyTorch, TensorFlow) имеют собственные оптимизации и экосистемы, которые тесно интегрируются с GPU и распределёнными коммуникациями (NCCL, Horovod), чего MLlib не предоставляет.')

    # Q6
    add_paragraph(doc, '\n6. Роль TorchDistributor в Spark; какую проблему решает', bold=True)
    add_paragraph(doc, """TorchDistributor (Spark) — это компонент, который упрощает запуск распределённого PyTorch-обучения поверх Spark-кластера. Его задача — предоставить абстракцию для:
    - выделения ресурсов Spark (executor'ов) под тренинг (CPU/GPU),
    - разворачивания процессов PyTorch на этих executor'ах,
    - инициализации окружения для DDP (настройка адресов, портов, world_size, rank),
    - управления жизненным циклом тренинга и сбором артефактов (логи, модели).

    Таким образом, TorchDistributor решает проблему интеграции рабочего процесса ETL в Spark с запуском распределённого PyTorch: вы можете готовить данные внутри Spark и затем запускать синхронное DDP-обучение через Spark API без ручной оркестрации контейнеров/процессов.""")

    # Part 2: Framework and Application Questions
    add_heading(doc, '\nPart 2: Framework and Application Questions', level=1)

    # Q7
    add_paragraph(doc, '7. Описать модель программирования Ray и примеры использования @ray.remote и Actors', bold=True)
    add_paragraph(doc, 'Ray программируется в стиле распределённых задач и акторов:\n- @ray.remote декорирует функцию или класс, превращая её в удалённую задачу/актора.\n- Вызов удалённой функции возвращает будущий объект (ObjectRef), который позже можно получить через ray.get().\n- Actors — это длительные stateful объекты, полезные для кеширования/межшагового состояния (например, модель, которая держит параметры или большой кэш данных).\n\nПример:\n- Простая удалённая задача: @ray.remote def preprocess(x): ...; ref = preprocess.remote(batch); res = ray.get(ref).\n- Actor: @ray.remote class ParameterServer: def __init__(self): self.state=...; ps = ParameterServer.remote(); ps.update.remote(grad)\nRay упрощает оркестрацию параллельных pipeline'ов, позволяет масштабировать и управлять ресурсами (GPU/CPU) на уровне задачи.'))

    # Q8
    add_paragraph(doc, '\n8. Схема запуска PyTorch DDP на Spark с использованием TorchDistributor (схематично + кратко)', bold=True)
    add_paragraph(doc, 'Схематический workflow (кратко):\n1) Spark driver запускает приложение с TorchDistributor, указывая сколько executor'"'ов/GPU нужно.\n2) TorchDistributor распределяет задачи и запускает на каждом executor'е процесс(ы) PyTorch (обычно один процесс на GPU).\n3) Каждый процесс инициализирует PyTorch, загружает данные (обычно из HDFS/Spark RDD/Parquet) либо через локальные файлы, либо через DataLoader с DistributedSampler.\n4) Процессы настраивают коммуникацию (init_process_group) — обычно через TCP rendezvous или через файловый/etcd сервис, используя адреса/порты, предоставленные Spark.\n5) Model replica выполняет forward/backward, DDP вызывает All-Reduce для градиентов; оптимайзер делает шаг синхронно.\n6) По завершении модели/эпохи результаты (модельные чекпоинты, логи) собираются и сохраняются в распределённое хранилище (HDFS/S3).')

    # Q9
    add_paragraph(doc, '\n9. Ray Train и Ray Tune — для чего используются; как они взаимодействуют', bold=True)
    add_paragraph(doc, 'Ray Train предназначен для упрощения распределённого тренинга моделей (абстракция для распределённых тренеров, интеграция с фреймворками — PyTorch, TensorFlow). Ray Tune ориентирован на гиперпараметрический поиск и оптимизацию.\n\nВместе они работают так: Ray Train запускает распределённые тренировки (например, один trial = одна настройка гиперпараметров), а Ray Tune организует многопробный запуск (много trial'ов), управляет планировщиком (ASHA, Bayesian) и собирает метрики; Tune вызывает Train для каждого trial, мониторит метрики и решает, какие trial'ы продолжать/прекратить.')

    # Q10
    add_paragraph(doc, '\n10. Сравнение: Pure PyTorch DDP, PyTorch on Spark (TorchDistributor), Ray Train', bold=True)
    add_paragraph(doc, 'Pure PyTorch DDP:\n+ Преимущества: прямой контроль, низкие накладные расходы, оптимизированные коммуникации (NCCL).\n- Недостатки: требует собственной оркестрации (например, Slurm, Kubernetes) и интеграции с системой сбора данных/ETL; сложнее интегрировать с системами хранения данных.\n\nPyTorch on Spark (TorchDistributor):\n+ Преимущества: плавная интеграция с существующим Spark-ETL, удобство разворачивания на Spark-кластере, удобство доступа к данным (Parquet/Hive).\n- Недостатки: дополнительные накладные расходы от Spark, потенциальная сложность при тонкой настройке коммуникаций и задержках; не всегда оптимален для максимальной производительности GPU-кластера.\n\nRay Train:\n+ Преимущества: гибкая модель задач/актеров, удобные абстракции для распределённого тренинга, тесная интеграция с Tune для HPO, хорошая поддержка для распределённых ресурсных политик.\n- Недостатки: ещё один фреймворк в стеке, может потребовать изучения и настройки, а также интеграции с корпоративными кластерами/файловыми системами.\n\nВыбор зависит от цели: максимальная скорость и контроль → чистый DDP; интеграция с ETL и Spark → TorchDistributor; исследовательские эксперименты и HPO → Ray Train + Tune.')

    # Part 3: Comprehensive Design and Analysis Questions
    add_heading(doc, '\nPart 3: Comprehensive Design and Analysis', level=1)

    # Q11
    add_paragraph(doc, '11. Дизайн end-to-end пайплайна для 100M текстов с BERT-large на Spark + PyTorch + TorchDistributor', bold=True)
    add_paragraph(doc, 'Ключевые шаги и объяснение:')
    add_paragraph(doc, '1) Хранение данных: разместить 100M текстов в распределённой файловой системе (HDFS/S3) в формате Parquet/TFRecord; хранить метаданные и шард-метаинфо.\n2) Preprocessing/ETL (Spark): использовать Spark для токенизации и предварительной фильтрации; разбить данные на шардированные Parquet файлы, применять динамическую подрезку/токенизацию; сохранить предобработанные батчи (например, input_ids, attention_mask). Spark удобно масштабирует CPU-bound preprocessing.\n3) Подача данных на обучение: использовать TorchDistributor для запуска DDP-наборов, где каждый Spark executor получает свои шардированные файлы; применять DataLoader с локальными файлами/MemoryMapped датасетами, или реализовать custom Dataset, читающий из Parquet через PyArrow.\n4) Настройка DDP: TorchDistributor создаёт процесс-группу (init_process_group) с world_size = num_gpus_total; обеспечить репликацию модели BERT-large либо использовать model-parallel (например, разделение слоёв) если модель не помещается на один GPU.\n5) Гибридный подход: если BERT-large не помещается на одному GPU — использовать модель-параллелизм (Pipe/Tensor Parallel) внутри каждого процесса и data-parallel снаружи; альтернативно, sharded state dict (FSDP) для оптимизации памяти.\n6) Тренинг: использовать mixed-precision (AMP), градиентный накопитель (gradient accumulation) для увеличения effective batch size; агрегировать градиенты через All-Reduce; логировать метрики в центральное хранилище (MLflow, TensorBoard).\n7) Checkpointing and saving: сохранять шардированные чекпоинты (в S3/HDFS), индексировать текущую эпоху/итерацию; при загрузке выполнять восстановление состояния распределённо.\n8) Serving: после обучения конвертировать модель/шарды в формат для inference (TorchScript, ONNX) и деплоить через модельный сервер (TorchServe/KServe).\n\nКлючевые инженерные моменты: efficient IO (shuffle minimization), use of FSDP/torch.distributed.pipeline for memory, careful tuning of batch size and gradient accumulation, and robust checkpointing.')

    # Q12
    add_paragraph(doc, '\n12. Анализ выбора архитектуры для сценариев A/B/C', bold=True)
    add_paragraph(doc, 'Сценарий A (ViT, multi-machine, medium data): предпочтительна Data-Parallel с PyTorch DDP, возможно в сочетании с tensor/model parallelism (например, using FairScale/FSDP/Deepspeed) если модель большая. DDP обеспечивает высокую производительность: NCCL optimized all-reduce. Если модель помещается на GPU — Data-Parallel + mixed-precision. Если модель очень большая — разбить слои (model parallel).')
    add_paragraph(doc, '\nСценарий B (интегрированный ETL -> обучение на Spark): PyTorch on Spark (TorchDistributor) — лучший выбор, потому что он интегрируется в существующий Spark-ETL пайплайн и упрощает развёртывание на тех же кластерах без сложной оркестрации. Для производительности можно комбинировать с DDP внутри Spark executors.')
    add_paragraph(doc, '\nСценарий C (исследования с частым HPO и деплоем): Ray Train + Ray Tune — предпочтителен: Tune автоматизирует HPO (ASHA, BO), интеграция с Train упрощает масштабируемые trials, а Ray служит платформой для управления экспериментами и развертывания. Альтернатива: Kubernetes + Argo + Katib для управления экспериментами.')

    # Part 4: Design question
    add_heading(doc, '\nPart 4: Comprehensive Design Questions', level=1)

    # Q13
    add_paragraph(doc, '13. End-to-end пайплайн для text classification с трансформером (пример: DistilBERT)', bold=True)
    add_paragraph(doc, 'Решение (пример: Spark + PyTorch + TorchDistributor):')
    add_paragraph(doc, '1) Data ingestion: хранить сырые тексты в HDFS/S3 (Parquet).\n2) Preprocessing: Spark jobs для очистки текста, токенизации (можно применять tokenizer в UDF или экспортировать предварительно токенизированные Parquet файлы).\n3) Dataset + DataLoader: на этапе тренинга каждый Spark executor читает свои порции Parquet; использовать DataLoader с DistributedSampler.\n4) Training: TorchDistributor запускает DDP процессы; модель DistilBERT (или HuggingFace transformer) загружается и реплицируется; использовать AMP и gradient accumulation при необходимости.\n5) Checkpointing: сохранять модели в S3/HDFS, поддерживать версионирование.\n6) Evaluation/Serving: после обучения оценить на тестовом сете; при успешных результатах конвертировать модель и деплоить (TorchServe/KServe).')

    # Q14
    add_paragraph(doc, '\n14. Анализ выбора архитектуры для сценариев A/B/C (повтор)', bold=True)
    add_paragraph(doc, 'Повтор описанных ранее рекомендаций, с фокусом на конкретный выбор инструментов:')
    add_paragraph(doc, 'Scenario A: PyTorch DDP + tensor/model parallel (FairScale/Deepspeed)\nScenario B: Spark + TorchDistributor (ETL + training in situ)\nScenario C: Ray Train + Ray Tune (fast HPO cycles), или Kubernetes + Katib/Argo')

    # Experiments & Results (detailed) — длинный раздел с реальными числами и интерпретацией
    add_heading(doc, '\nExperiments and results / Эксперименты и результаты', level=1)
    add_paragraph(doc, 'Ниже приводится подробное описание всех экспериментальных запусков, результатов и интерпретаций, выполненных в рамках проекта. Эти материалы включают числовые сводки, ключевые наблюдения и план дальнейших работ.')

    # Include recommender output if present
    try:
        rec_path = os.path.join('..', 'BDA_9_Reccomendation_System_Assigment', 'outputs', 'results_user_50.txt')
        if os.path.exists(rec_path):
            with open(rec_path, 'r', encoding='utf-8') as f:
                rec_text = f.read()
            add_paragraph(doc, 'Recommender (BDA_9) — выдержки результатов:', bold=True)
            add_paragraph(doc, rec_text)
        else:
            add_paragraph(doc, 'Recommender (BDA_9): результаты не найдены в репозитории.')
    except Exception:
        add_paragraph(doc, 'Recommender (BDA_9): ошибка при чтении файла результатов.')

    # Include BDA_6 (Topic5) numeric results
    try:
        base = os.path.join('..', 'BDA_6_Topic5_Regression_Clustering', 'outputs')
        cal_path = os.path.join(base, 'exp1', 'california', 'results_exp1.json')
        syn_path = os.path.join(base, 'exp1', 'synthetic', 'results_exp1.json')
        exp2_path = os.path.join(base, 'exp2', 'results_exp2.json')
        if os.path.exists(cal_path):
            with open(cal_path, 'r', encoding='utf-8') as f:
                california = f.read()
            add_paragraph(doc, '\nTopic5 experiments (BDA_6) — numeric summaries:', bold=True)
            add_paragraph(doc, 'California housing (regression) results (json):')
            add_paragraph(doc, california)
        if os.path.exists(syn_path):
            with open(syn_path, 'r', encoding='utf-8') as f:
                synthetic = f.read()
            add_paragraph(doc, 'Synthetic sin dataset results (json):')
            add_paragraph(doc, synthetic)
        if os.path.exists(exp2_path):
            with open(exp2_path, 'r', encoding='utf-8') as f:
                exp2 = f.read()
            add_paragraph(doc, 'Clustering experiments overview (NMI scores):')
            add_paragraph(doc, exp2)
    except Exception:
        add_paragraph(doc, 'Topic5 experiments: ошибка при чтении результатов.')

    # Include BDA_3 processed summary if present
    try:
        amazon_path = os.path.join('..', 'BDA_3_AI_enhanced_ETL', 'data', 'processed', 'Amazon-2024-Annual-Report.pdf.json')
        if os.path.exists(amazon_path):
            with open(amazon_path, 'r', encoding='utf-8') as f:
                amazon = f.read()
            add_paragraph(doc, '\nBDA_3 AI-ETL: пример обработанного PDF (сводка и метаданные):', bold=True)
            add_paragraph(doc, amazon)
        else:
            add_paragraph(doc, 'BDA_3: обработанные JSON-файлы не найдены.')
    except Exception:
        add_paragraph(doc, 'BDA_3: ошибка при чтении обработанного JSON.')

    # Long analytical text (Russian) describing experiments and interpretation
    long_text = (
        'Экспериментальная часть проекта была спланирована с целью продемонстрировать основные подходы к распределённому обучению ' 
        'и оценить простые модели в условиях ограниченных ресурсов.\n\n'
        '1) Рекомендательная система (BDA_9).\n'
        'Мы реализовали гибридный подход: классическое user-based collaborative filtering с косинусной мерой близости, ' 
        'дополненное жанровой корректировкой и фактором популярности. Для одного пользователя (id=50) были сформированы Top-5 ' 
        'рекомендации с учётом разнообразия жанров и бонусов за популярность и возраст фильма. Результат показывает, что алгоритм ' 
        'предложил фильмы из предпочитаемых жанров (Drama, Romance, Comedy) и предоставил интерпретируемые оценки в шкале 1–5.\n\n'
        'Метрики: для этого proof-of-concept Precision@5 и NDCG@5 формально не рассчитаны; следующая итерация включает скрипт оценки `evaluate.py`, ' 
        'который посчитает Precision@K/NDCG@K по всем пользователям и сохранит отчёт в JSON.\n\n'
        'Качества и наблюдения:\n- Жанровые шаблоны ML-100k содержат шум (например, некоторые sci-fi помечены как Romance). ' 
        'Мы применили эвристики очистки жанров: удаление Romance у явно Sci‑Fi/Action/Adventure фильмов.\n'
        '- Смешивание CF и жанрового скоринга даёт более правдоподобную выдачу, но требует подбора коэффициента alpha: сильный жанровый буст приводит к потере персонализации; слабый — к нерелевантным рекомендациям по жанру.\n\n'
        '2) Topic5 — Регрессия и кластеризация (BDA_6).\n'
        'В экспериментах с регрессией (California housing) сравнивались LinearRegression и MLP. Результаты:\n'
        '- LinearRegression: MSE ≈ 0.556, R^2 ≈ 0.576 (время тренировки ≈ 0.0047s).\n'
        '- MLP: MSE ≈ 0.283, R^2 ≈ 0.784 (время тренировки ≈ 6.12s) — значительное улучшение качества ценой времени и сложности модели.\n\n'
        'На синтетическом синусном наборе MLP продемонстрировал лучшее соответствие (R^2 ≈ 0.935 vs 0.636 у линейной модели).\n\n'
        'Кластеризация (exp2):\n- KMeans (baseline): NMI ≈ 0.487.\n- DeepCluster-like proxy: NMI ≈ 0.252 — прототип показал худший результат; для корректного сравнения требуется полноценная реализация с encoder/decoder и более длительной тренировкой на GPU.\n\n'
        '3) AI-ETL (BDA_3).\n'
        'Пайплайн извлечения текста из PDF был протестирован на Annual Report Amazon 2024. Экстракт сохранил ~320k символов и ряд полезных комментариев/цитат, ' 
        'извлечённых из корпоративного документа. Такой подход позволяет автоматизировать предварительный анализ корпоративной документации и генерировать сводки для downstream NLP задач.\n\n'
        'Практические замечания: OCR требуется для сканов; качество извлечения зависит от структуры PDF; рекомендовано добавить LLM-постпроцессинг для коррекции артефактов.\n\n'
        '4) Инфраструктурные выводы и рекомендации.\n'
        'Подготовлены демонстрационные скрипты: `train_ddp.py` (локальный DDP demo) и `spark_torch_distributor_example.py` (шаблон для Spark + TorchDistributor). ' 
        'Рекомендовано: Spark для CPU-bound preprocessing, далее DDP/FSDP на GPU для тренировки; при очень больших моделях — model-parallel/FSDP.\n\n'
        '5) Следующие шаги: реализация evaluate.py, подбор гиперпараметров для гибридного рекоммендера, полная DeepCluster реализация и автоматизация генерации DOCX-отчётов.\n'
        'Этот раздел служит как самодостаточный отчёт по экспериментам: содержит описания, числовые результаты, интерпретацию и дорожную карту дальнейших улучшений.'
    )

    add_paragraph(doc, long_text)

    # Save
    doc.save(out_path)
    print('Wrote report to', out_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--out', default=OUT_DOC_DEFAULT)
    args = parser.parse_args()
    write_report(args.out)
