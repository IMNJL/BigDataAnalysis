"""
Generate DOCX report for Topic 5 experiments using experiment outputs JSONs.
"""
import argparse
import json
import os
from docx import Document
from docx.shared import Pt, Inches


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def insert_architecture(doc, screenshots_dir='screenshots'):
    add_heading(doc, 'Part 2: Architecture (Diagram)', level=2)
    arch = os.path.join(screenshots_dir, 'architecture.png')
    if os.path.exists(arch):
        try:
            doc.add_picture(arch, width=Inches(6))
            add_paragraph(doc, 'Figure: Architecture diagram')
        except Exception as e:
            add_paragraph(doc, f'Failed to insert architecture diagram: {e}')
    else:
        add_paragraph(doc, 'Architecture diagram not found. Place screenshots/architecture.png to include it.')


def insert_steps(doc, steps_path='steps.json', screenshots_dir='screenshots'):
    add_heading(doc, 'Part 3: Steps and Screenshots', level=2)
    if os.path.exists(steps_path):
        try:
            with open(steps_path, 'r', encoding='utf-8') as f:
                steps = json.load(f)
        except Exception as e:
            add_paragraph(doc, f'Failed to read steps file: {e}')
            steps = []
    else:
        steps = []

    if not steps:
        add_paragraph(doc, 'No steps.json found. Example format:')
        add_paragraph(doc, '{"steps": [{"title":"Step 1: Crawl","prompt":"<prompt>","response":"<response>","screenshot":"step1.png"}, ...]}')
        return

    captions = {}
    captions_path = os.path.join(screenshots_dir, 'captions.json')
    if os.path.exists(captions_path):
        try:
            with open(captions_path, 'r', encoding='utf-8') as cf:
                captions = json.load(cf)
        except Exception:
            captions = {}

    for s in steps.get('steps', []):
        title = s.get('title') or 'Step'
        add_paragraph(doc, title)
        prompt = s.get('prompt', '')
        resp = s.get('response', '')
        if prompt:
            add_paragraph(doc, 'Prompt:')
            add_paragraph(doc, prompt)
        if resp:
            add_paragraph(doc, 'Response:')
            add_paragraph(doc, resp)
        img = s.get('screenshot')
        if img:
            img_path = os.path.join(screenshots_dir, img)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    cap = captions.get(img) if isinstance(captions, dict) else None
                    if cap:
                        add_paragraph(doc, f'Caption: {cap}')
                except Exception as e:
                    add_paragraph(doc, f'Failed to insert screenshot {img}: {e}')


def insert_final_result(doc, exp1_path=None, exp2_path=None):
    add_heading(doc, 'Final Result', level=2)
    add_paragraph(doc, 'Summary of main findings and conclusions.')
    if exp1_path and os.path.exists(exp1_path):
        try:
            with open(exp1_path, 'r', encoding='utf-8') as f:
                j = json.load(f)
            add_paragraph(doc, 'Experiment 1 overview:')
            add_paragraph(doc, json.dumps(j, ensure_ascii=False, indent=2)[:2000])
        except Exception as e:
            add_paragraph(doc, f'Failed to read exp1 results: {e}')
    if exp2_path and os.path.exists(exp2_path):
        try:
            with open(exp2_path, 'r', encoding='utf-8') as f:
                j = json.load(f)
            add_paragraph(doc, 'Experiment 2 overview:')
            add_paragraph(doc, json.dumps(j, ensure_ascii=False, indent=2)[:2000])
        except Exception as e:
            add_paragraph(doc, f'Failed to read exp2 results: {e}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--steps', help='path to steps.json (prompt/response pairs)', default='steps.json')
    parser.add_argument('--screenshots', help='screenshots directory', default='screenshots')
    parser.add_argument('--exp1', help='path to exp1 results JSON', default='outputs/exp1/results_exp1_overview.json')
    parser.add_argument('--exp2', help='path to exp2 results JSON', default='outputs/exp2/results_exp2.json')
    parser.add_argument('--out', help='output docx path', default='report_topic5.docx')
    args = parser.parse_args()

    doc = Document()
    doc.add_heading('BDA Topic 5 - Regression and Clustering Experiments')

    # Part 1: Experiment Plan
    add_heading(doc, 'Part 1: Experiment Plan', level=2)
    add_paragraph(doc, 'Goals:')
    add_paragraph(doc, '- Compare regression methods (Linear, MLP) under controlled and real-world datasets.')
    add_paragraph(doc, "- Compare clustering: KMeans vs DeepCluster-like on image datasets (MNIST, Fashion-MNIST).")
    add_paragraph(doc, '- Metrics: MSE, R2, training time for regression; NMI for clustering; qualitative analysis and interpretability.')

    # Part 2: Architecture
    insert_architecture(doc, args.screenshots)

    # Part 3: Steps and Screenshots
    insert_steps(doc, args.steps, args.screenshots)

    # Final Result
    insert_final_result(doc, args.exp1, args.exp2)

    doc.save(args.out)
    print('Wrote report to', args.out)
