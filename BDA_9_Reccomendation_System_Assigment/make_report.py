# ...existing code...
import os
import argparse
from docx import Document
from docx.shared import Pt
import datetime

PROJECT_ROOT = os.path.dirname(__file__)
OUTPUTS = os.path.join(PROJECT_ROOT, "outputs")
REPORT_PATH = os.path.join(PROJECT_ROOT, "BDA_9_Recommendation_report.docx")


def _safe_read_result(user_id):
    result_file = os.path.join(OUTPUTS, f"results_user_{user_id}.txt")
    if os.path.exists(result_file):
        with open(result_file, "r", encoding="utf-8") as f:
            return f.read()
    return None


def _add_paragraphs_from_text(doc, text):
    # Split by double newlines to get paragraphs, keep single-line breaks as within-paragraph formatting
    for para in [p.strip() for p in text.split('\n\n') if p.strip()]:
        p = doc.add_paragraph(para)


def build_report(user_id):
    """
    Build a comprehensive recommendation assignment report containing:
      - Part 1: Theoretical Analysis (comparative table + similarity metrics + SVD)
      - Part 2: Algorithm implementation summary and LLM tag generation
      - Part 3: Case design for Taobao (cold-start + interpretability)
      - Part 4: Frontier thinking questions and answers
      - Bonus: End-to-end LLM movie recommender discussion

    The final .docx is written to REPORT_PATH.
    """
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading('BDA_9 Recommendation — Comparative Study and Optimization', level=1)
    doc.add_paragraph(f'Generated: {now}')

    # Intro
    doc.add_heading('Executive summary', level=2)
    exec_summary = (
        "This document contains Parts 1 through 4 of the assignment plus an optional Bonus section. "
        "It compares traditional Collaborative Filtering (CF) with LLM-enhanced recommendation systems, "
        "explains similarity measures and matrix factorization, describes an implemented User-Based CF method, "
        "illustrates how an LLM can generate user profile tags, designs a Taobao cold-start and interpretability scheme, "
        "and discusses frontier research questions. The content is prepared for submission and reproducibility."
    )
    doc.add_paragraph(exec_summary)

    # --- Part 1
    doc.add_heading('Part 1: Theoretical Analysis', level=2)

    doc.add_heading('1. Comparative Analysis: CF vs LLM-enhanced systems', level=3)
    comp_text = (
        "This subsection compares Traditional Collaborative Filtering (CF) and LLM-enhanced recommendation systems across four aspects: "
        "Data Utilization Methods, Recommendation Logic, Interpretability, and the Cold-Start Problem. Each aspect is discussed below in detail."
    )
    doc.add_paragraph(comp_text)

    # Data Utilization
    doc.add_heading('Data Utilization Methods', level=4)
    du = (
        "Traditional CF: Primarily uses user-item interaction signals such as explicit ratings, likes, purchases, and implicit signals (views, clicks). "
        "These are represented as a sparse user-item matrix or as edges in a bipartite interaction graph. CF relies on the statistical aggregation of these interactions to infer similarity and preference. "
        "Data modalities are limited unless the system is explicitly hybridized with content features.\n\n"
        "LLM-Enhanced: Embraces heterogeneous textual and multimodal sources. Item metadata (titles, descriptions), user-generated content (reviews, comments), session logs and even free-text user responses are converted into embeddings or structured tags. "
        "LLMs can transform sparse signals into dense semantic features and can integrate cross-domain textual knowledge not present in raw interaction matrices."
    )
    _add_paragraphs_from_text(doc, du)

    # Recommendation Logic
    doc.add_heading('Recommendation Logic', level=4)
    rl = (
        "Traditional CF: Operates with neighborhood-based methods (user-based or item-based) or model-based approaches (matrix factorization, SVD, probabilistic factor models). "
        "Neighborhood methods find similar users or items with explicit similarity metrics and aggregate known preferences. Model-based methods learn latent factors that explain the interaction patterns.\n\n"
        "LLM-Enhanced: Relies on semantic similarity and natural language reasoning. Approaches include generating embeddings for semantic retrieval, producing semantic tags for downstream models, or performing end-to-end generation (directly outputting ranked candidates or natural-language recommendations). "
        "Because LLMs encode knowledge from large corpora, they can generalize across sparse or unseen items and answer open-ended queries."
    )
    _add_paragraphs_from_text(doc, rl)

    # Interpretability
    doc.add_heading('Interpretability', level=4)
    interp = (
        "Traditional CF: Often provides transparent explanations. Neighborhood methods can produce explanations such as 'Users similar to you liked X' or 'Because you liked Y, you may like Z'. Latent-factor explanations can be approximated by mapping dimensions to human concepts but are less direct.\n\n"
        "LLM-Enhanced: Generates fluent, human-like explanations and can combine multiple signals into one sentence (e.g., 'We recommend X because you like sci-fi and rated Y highly'). However, explanations from LLMs may be only plausible and not causally tied to the internal decision process unless the system enforces grounding and provenance constraints."
    )
    _add_paragraphs_from_text(doc, interp)

    # Cold-start
    doc.add_heading('Cold-Start Problem', level=4)
    cs = (
        "Traditional CF: Poor performance when users or items have no or very few interactions. Cold-start solutions typically include asking onboarding questions, using content-based techniques, or leveraging side information (demographics, item attributes).\n\n"
        "LLM-Enhanced: Better positioned for cold-start since LLMs can infer preferences from short textual inputs (user descriptions, browsing snippets), map free-text to taxonomy tags, and produce embeddings that link new items to existing semantic neighborhoods. LLM prompting and few-shot learning can synthesize a profile with minimal explicit interaction."
    )
    _add_paragraphs_from_text(doc, cs)

    # Table-like summary (textual)
    doc.add_heading('Summary: Side-by-side', level=4)
    summary = (
        "- Data: CF = interaction matrices; LLM = text + embeddings + metadata.\n"
        "- Logic: CF = similarity/latent factors; LLM = semantic inference and generation.\n"
        "- Interpretability: CF = traceable neighborhood signals; LLM = fluent but potentially ungrounded narratives.\n"
        "- Cold-start: CF = limited; LLM = strong with textual/metadata features."
    )
    doc.add_paragraph(summary)

    doc.add_heading('How the three LLM paradigms compensate for CF shortcomings', level=3)
    paradigms = (
        "Three LLM paradigms are commonly used: (1) Embedding Generation, (2) Semantic Tag Generation, (3) End-to-End Generation. Each addresses CF weaknesses as follows.\n\n"
        "1) Embedding Generation: LLMs or text encoders convert item descriptions, user reviews and even short user bios into dense vectors. These dense vectors allow semantic nearest-neighbor retrieval even when explicit co-occurrence is rare. Embeddings make it possible to compute similarity between newly released items and users without collaborative history.\n\n"
        "2) Semantic Tag Generation: LLMs extract concise semantic labels or tags (genres, themes, moods) from small text or sparse interactions. These tags act as dense side-features for classical recommenders or as inputs to content-based matching, bridging the gap when collaborative signals are absent.\n\n"
        "3) End-to-End Generation: The LLM directly proposes ranked lists, explanations, or personalized text-based suggestions by reasoning over provided context. This is valuable for zero-shot scenarios, conversational recommenders, and when integrating knowledge beyond the dataset. The main trade-offs are groundedness and evaluation: output must be validated and controlled to avoid hallucination."
    )
    _add_paragraphs_from_text(doc, paradigms)

    # --- Part 1.2 Similarity and SVD
    doc.add_heading('Similarity Calculation and Matrix Factorization', level=3)

    doc.add_heading('Similarity measures: Jaccard, Cosine, Pearson', level=4)
    sim_text = (
        "This section lists the formulas, appropriate scenarios, and limitations for three commonly used similarity measures.\n\n"
        "Jaccard similarity (for sets): J(A,B) = |A ∩ B| / |A ∪ B|. Use when interactions are binary (purchased / not purchased, clicked / not clicked). Advantages: interpretable for set overlap; robust when only presence/absence matters. Limitations: ignores frequency and rating magnitudes; can be misleading when item sets are very small.\n\n"
        "Cosine similarity (for vectors): cos(u,v) = (u · v) / (||u|| ||v||). Use for numerical interaction vectors (ratings, tf-idf counts, embeddings). Advantages: scale-invariant, measures orientation of preference vectors, widely used for sparse numeric data. Limitations: does not correct for user rating bias and can be influenced by shared zeros in extremely sparse matrices.\n\n"
        "Pearson correlation (centered measure): r_{uv} = Σ_i (r_{ui}-μ_u)(r_{vi}-μ_v) / (sqrt(Σ_i (r_{ui}-μ_u)^2) sqrt(Σ_i (r_{vi}-μ_v)^2)). Use when users have different rating baselines to capture linear correlation in co-rated items. Advantages: mitigates differences in average rating level across users. Limitations: unreliable when the number of co-rated items is small; sensitive to outliers."
    )
    _add_paragraphs_from_text(doc, sim_text)

    doc.add_heading('Role of SVD / Matrix Factorization', level=4)
    svd_text = (
        "Matrix factorization techniques (SVD, ALS, probabilistic MF) decompose the user-item matrix into low-dimensional latent factors that capture patterns of co-preference. If R is the user-item matrix, SVD approximates R ≈ U Σ V^T where U and V contain user/item latent vectors.\n\n"
        "How SVD helps sparsity: By learning shared latent factors it generalizes across users and items, enabling prediction for unobserved user-item pairs. Regularization and low-rank constraints reduce overfitting and denoise sparse interactions. Practically, implementations use SGD or alternating least squares (ALS) on the observed entries. SVD is a core building block of many scalable recommenders."
    )
    _add_paragraphs_from_text(doc, svd_text)

    # --- Part 2 Implementation
    doc.add_heading('Part 2: Algorithm Implementation & Programming', level=2)

    doc.add_heading('1. User-Based Collaborative Filtering (implementation summary)', level=3)
    part2_text = (
        "Implementation choices and rationales:\n"
        "- Dataset: MovieLens-100K (standard benchmark).\n"
        "- Similarity: Cosine similarity between full user rating vectors (missing ratings filled as zero in this simple implementation). Cosine is chosen because it naturally measures similarity in preference patterns independent of magnitude.\n"
        "- Prediction: For a target user, select top-K neighbors by cosine similarity, compute weighted average of neighbor ratings (weights = similarity), and recommend the top-N items the target has not rated.\n\n"
        "Notes on the simple design: Filling missing entries with zeros is a pragmatic choice for cosine computation but is not ideal in production: one should compute similarities only over co-rated items or use adjusted-cosine / Pearson to handle user biases. The provided code in the project follows the basic approach for clarity and reproducibility."
    )
    _add_paragraphs_from_text(doc, part2_text)

    doc.add_heading('Example output and reasoning', level=4)
    reasoning = (
        "The Top-5 recommendations returned by the script are items that similar users rated highly. The reasoning behind each recommendation can be explained with: 'users with similar rating patterns enjoyed this movie', and optionally by mapping neighbor ratings to tags or genres for user-friendly explanations."
    )
    doc.add_paragraph(reasoning)

    doc.add_heading('2. LLM-generated user profile tags', level=3)
    llm_section = (
        "Approach: Use an open-source instruction-following LLM (for example, FLAN-T5 or similar) to convert short user histories into compact tags. The pipeline uses few-shot prompting with examples mapping short histories to comma-separated tags.\n\n"
        "Example prompt structure (few-shot):\n"
        "Example 1: Input: Rated 'The Matrix' 5, 'Inception' 5, 'Titanic' 2 -> Output: sci-fi, action, psychological thriller.\n"
        "Given a new user history, the model outputs tags such as 'sci-fi, drama, romance', which can be used as features in a hybrid recommender or to initialize a cold-start profile.\n\n"
        "Practical considerations: choose a compact model for CPU usage (e.g., flan-t5-small) for local experimentation; for production or higher accuracy, larger models or fine-tuned variants can be used."
    )
    _add_paragraphs_from_text(doc, llm_section)

    # Include results file content if present
    doc.add_heading('Included recommendation output (if available)', level=3)
    results = _safe_read_result(user_id)
    if results:
        doc.add_paragraph(f'Recommendation output for user {user_id}:')
        doc.add_paragraph(results)
    else:
        doc.add_paragraph(f'No precomputed recommendation output found for user {user_id}. Run run/recommend_user.py to produce outputs/results_user_{user_id}.txt and re-run this script.')

    # --- Part 3 Case Design
    doc.add_heading('Part 3: Case Design and Analysis (Taobao)', level=2)

    doc.add_heading('1. Cold-Start Scenario Design for new users', level=3)
    cold_start = (
        "Design goal: Provide useful recommendations for brand-new users with zero or minimal interaction history by leveraging LLMs and available side-channel signals.\n\n"
        "Pipeline: \n"
        "1) Onboarding prompt + structured inputs: gather quick multi-choice preferences and a single free-text box (e.g., 'I like sportswear and gadgets').\n"
        "2) Prompting + Semantic Tag Generation: send the onboarding text and short session logs to an LLM prompt which outputs standardized tags mapped to internal taxonomy.\n"
        "3) Instruction Tuning: use a small instruction-tuned model that enforces tag taxonomy and reduces hallucination; this model is trained to consistently map free text to the internal tag set.\n"
        "4) Candidate generation: convert tags to an embedding vector and retrieve nearest items in product embedding index. Optionally, augment with collaborative signals from cohort-level behavior (popularity among similar demographic segments).\n"
        "5) Rapid feedback loop: show a curated mix of confident matches and exploratory items; capture first-session clicks to refine the profile quickly.\n\n"
        "Data utility: registration fields (age, gender, location), social graph (friends' public favorites), and browsing logs (dwell time, categories visited) are transformed into prompts or features for the LLM. Social signals can be incorporated by aggregating tags from connected accounts to bootstrap recommendations."
    )
    _add_paragraphs_from_text(doc, cold_start)

    doc.add_heading('2. Interpretability Enhancement Scheme', level=3)
    interp_scheme = (
        "Module design: A constrained LLM explanation module produces personalized, evidence-backed explanations. The module receives: user tags, recent high-rated items, item metadata, and a small list of provenance facts. The prompt instructs the LLM to generate short sentences that must cite at least one user-specific fact (e.g., 'You rated X 5/5') or an explicit tag.\n\n"
        "Advantages of LLM explanations: personalization, fluency, ability to combine multiple signals (ratings, tags, recency) into a natural phrase. Disadvantages: risk of hallucination if provenance is not strictly enforced; higher compute cost.\n\n"
        "Comparison to CF explanations: CF explanations are easier to verify ('Users similar to you also liked X') and are intrinsically traceable to neighbor statistics, while LLM explanations are more persuasive and user-friendly but require strict grounding to maintain veracity. A hybrid approach (use CF provenance + LLM surface text) often delivers the best balance."
    )
    _add_paragraphs_from_text(doc, interp_scheme)

    # --- Part 4 Frontier Thinking
    doc.add_heading('Part 4: Frontier Thinking Questions', level=2)

    q1 = (
        "Question: Can LLMs completely replace traditional Collaborative Filtering?\n\n"
        "Answer: In my assessment, a complete replacement is unlikely in the near-to-mid term. LLMs bring exceptional semantic understanding and the ability to infer preferences from sparse textual inputs; however, collaborative filtering captures aggregate behavioral co-occurrence patterns at scale which are directly predictive of future interactions. CF benefits from efficient, well-understood, and scalable optimization methods for personal ranking (matrix factorization, approximate nearest neighbors over latent vectors). The best practical architecture is hybrid: use CF for scalable personalized ranking and LLMs to augment features, handle cold-start, and provide natural-language interfaces."
    )
    _add_paragraphs_from_text(doc, q1)

    q2 = (
        "Question: Which will be the dominant trend — LLM as main component or CF as main component with LLM-enhancement?\n\n"
        "Answer: Likely CF + hybrid models will remain the backbone for large-scale ranking with LLMs as strong enhancers. CF is data-efficient for users with rich history and supports tight evaluation metrics (A/B tests, offline ranking losses). LLMs will be increasingly integrated for representation learning, candidate generation, user-facing explanation, and personalization in sparse regimes. Over time, architectures may become more LLM-centric for small- to medium-scale systems or for specialized experiences (conversational recommenders), but at web scale CF-based pipelines remain core."
    )
    _add_paragraphs_from_text(doc, q2)

    # --- Bonus
    doc.add_heading('Bonus: End-to-End LLM Movie Recommender (design)', level=2)
    bonus = (
        "Task: Build an end-to-end LLM recommender where the input is a textual history and the output is a ranked list with reasons.\n\n"
        "Design notes: Use a two-stage approach to reduce hallucination and enforce quality: 1) Candidate generation by semantic retrieval (embed user text and retrieve a candidate set from an item corpus), 2) LLM re-ranking and explanation, where the LLM receives the candidate list plus short provenance (titles + metadata) and user text and outputs a final ranked list with concise rationales.\n\n"
        "Implementation tip: For local experimentation, generate tags with a small instruction-following LLM and then use a precomputed item metadata table to retrieve candidates. The re-ranker prompt should be constrained and include the candidate set to avoid inventing unknown titles. This hybrid design reduces hallucination while maintaining the flexibility of LLM natural language."
    )
    _add_paragraphs_from_text(doc, bonus)

    # Appendix: references and run instructions
    doc.add_heading('Appendix: How to reproduce', level=3)
    run_instructions = (
        "1) Ensure MovieLens-100K is present at data/raw/ml-100k.\n"
        "2) From the project root run (macOS, zsh):\n\n"
        "   python3 make_report.py --user 50\n\n"
        "This will produce 'BDA_9_Recommendation_report.docx' at the project root. If recommendation outputs exist in 'outputs/results_user_{user}.txt', they will be included in the report.\n\n"
        "Notes: The script uses python-docx; to install dependencies: pip install python-docx pandas scikit-learn transformers torch (depending on which modules you plan to run)."
    )
    _add_paragraphs_from_text(doc, run_instructions)

    # Save document
    doc.save(REPORT_PATH)
    print(f'Report written to: {REPORT_PATH}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Build automated recommendation report')
    parser.add_argument('--user', type=int, default=50, help='User id to include outputs for')
    args = parser.parse_args()
    build_report(args.user)
# ...existing code...